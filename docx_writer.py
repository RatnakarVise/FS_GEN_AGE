from docx import Document
from docx.shared import Pt, RGBColor
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
import re

def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.size = Pt(14)

def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        start, end = match.span()
        paragraph.add_run(text[cursor:start])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = end
    paragraph.add_run(text[cursor:])

def add_code_block(doc, code_lines):
    para = doc.add_paragraph()
    run = para.add_run("\n".join(code_lines))
    # run.font.name = "Courier New"
    run.font.size = Pt(10)

def add_markdown_table(doc, lines):
    if len(lines) < 2:
        return
    # Remove markdown separator like |----|
    header_line = lines[0]
    separator_line = lines[1] if len(lines) > 1 else ""

    headers = [h.strip() for h in header_line.strip('|').split('|')]
    if not headers:
        return
    # Only accept second line if it contains dashes (markdown separator)
    if not re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", separator_line):
        rows = lines[1:]
    else:
        rows = lines[2:]

    doc_table = doc.add_table(rows=1, cols=len(headers))
    doc_table.style = 'Table Grid'

    # Header
    for i, h in enumerate(headers):
        doc_table.rows[0].cells[i].text = h

    # Rows
    for row_line in rows:
        if not row_line.strip():
            continue
        row_cells = [cell.strip() for cell in row_line.strip('|').split('|')]
        if len(row_cells) != len(headers):
            continue  # skip malformed rows
        row = doc_table.add_row().cells
        for i, cell in enumerate(row_cells):
            row[i].text = cell  

def create_docx(ts_text: str, buffer):
    doc = Document() 
    doc.add_heading('FUNCTIONAL SPECIFICATION', level=1)

    lines = ts_text.splitlines()
    current_section = ""
    current_content = []
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    current_paragraph_lines = []
    #section_header_pattern= re.compile(r"^\s*(\d{1,2})\.\s"(.?):\s(.+)?$")
    #plain_header_pattern = re.compile(r"^\s*(\d{1,2})\. \s*(.+)$")
    #subheading pattern = re.compile(r"^\s*(\d{1,2})\.(\d+)\s*(.+?):25")

    table_line_pattern = re.compile(r"^\|(.+?)\|$")
    def flush_paragraphs():
        for line in current_paragraph_lines:
            add_paragraph(doc, line)
        current_paragraph_lines.clear()

    def flush_table():
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines.clear()

    def flush_current_content():
        if current_section:
            add_heading(doc, current_section)
        for para in current_content:
            add_paragraph(doc, para)

    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Handle code block
        if line.startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                flush_paragraphs()
                add_code_block(doc, code_block_lines)
                code_block_lines.clear()
            continue
        if in_code_block:
            code_block_lines.append(line)
            continue
        if  table_line_pattern.match(line):
            table_lines.append(line)
            in_table = True
            continue
        elif in_table:
            flush_paragraphs()
            flush_table()
            in_table = False
        current_paragraph_lines.append(line)

    # Final flush
    flush_paragraphs()
    flush_table()

    doc.save(buffer)